# -*- coding: utf-8 -*-
"""
Created on Wed May  3 16:37:09 2023
"""

import streamlit as st
from PyPDF2 import PdfReader
import docx
import os

# Streamlit app title and description
st.title("PDF and Word Comparison")
st.write("Select the PDF file and Word document to compare and save the differences.")

# Choose the PDF file
pdf_file = st.file_uploader("Select PDF File", type="pdf")
if pdf_file is not None:
    reader = PdfReader(pdf_file)

    # Choose the Word document
    word_file = st.file_uploader("Select Word Document", type="docx")
    if word_file is not None:
        doc = docx.Document(word_file)

        # List to store extracted text from all pages
        pages_text = []

        # Iterating over all pages and extracting text from PDF
        for page in range(len(reader.pages)-1):
            current_page = reader.pages[page]
            text = current_page.extract_text()
            pages_text.append(text)

        all_text = "\n".join(pages_text)

        # Get the text from the Word document
        doc_text = '\n'.join([para.text for para in doc.paragraphs])

        # Split the texts into words
        words1 = doc_text.split()
        words2 = all_text.split()

        # Get the differences between the words
        differences = set(words2) - set(words1)

        # Print the differences
        st.subheader("Differences:")
        differing_words = list(differences)

        for word in differing_words:
            st.write(f'- {word}')

        # Save differing words in a Word file
        output_file = st.text_input("Output file name", value="Differences.docx")
        output_folder_path = st.text_input("Output folder path", value="path/to/folder")

        if st.button("Save Differences"):
            if output_file:
                output_filename = output_file.strip()
                output_path = os.path.join(output_folder_path, output_filename)

                # Create the Word document for output
                output_doc = docx.Document()
                output_doc.add_paragraph("Differences between PDF and Word document:")

                # Create a bulleted list
                list_paragraph = output_doc.add_paragraph()
                list_paragraph.style = output_doc.styles["List Bullet"]

                for word in differing_words:
                    list_paragraph.add_run(f"{word}\n")

                output_doc.save(output_path)
                st.success(f"Differences saved to: {output_path}")
            else:
                st.error("Please enter a valid output file name.")
